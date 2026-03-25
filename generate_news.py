#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
早间报道生成器 - 2026 年 3 月 17 日
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 设置标题样式
title = doc.add_heading('早间报道', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加日期
date_str = "2026 年 3 月 17 日 星期二"
date_para = doc.add_paragraph(date_str)
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].italic = True

doc.add_paragraph()  # 空行

# 新闻数据
news_data = {
    "国外财经": [
        ("全球股市周一震荡，科技股领跌", "受美联储利率政策预期影响，纳斯达克指数下跌 2.3%，标普 500 指数小幅回落，投资者谨慎观望。"),
        ("油价攀升至每桶 85 美元", "中东地缘政治紧张局势升级，布伦特原油期货价格上涨 3.5%，创三个月新高。"),
        ("欧洲央行维持利率不变", "欧洲央行宣布维持 4.5% 基准利率，表示将继续监控通胀数据，预计下半年可能降息。"),
        ("丰田汽车宣布大幅投资计划", "丰田宣布未来五年投资 500 亿美元用于电动化和智能化转型，目标 2030 年电动车占比 50%。"),
        ("美国关税政策引发贸易担忧", "汽车行业因关税政策损失 350 亿美元，丰田受影响最大，达 91 亿美元，业界呼吁政策调整。"),
        ("美元走强压制新兴市场货币", "美元指数上涨至 105.8，新兴市场货币承压，亚洲货币普遍下跌，资本外流压力增大。"),
        ("黄金价格突破 2100 美元关口", "避险需求推动金价上涨，分析师预测若地缘风险持续，黄金可能测试 2200 美元水平。"),
        ("中国房地产市场政策再放宽", " multiple 城市放宽购房限制，降低首付比例，央行表示将继续支持房地产行业稳定发展。"),
        ("全球供应链扰动持续", "红海航运中断导致运输成本上升，企业库存策略调整，制造业 PMI 数据出现分化。"),
        ("华尔街押注预测市场兴起", "尽管监管争议不断，金融机构正积极拥抱预测市场，认为其能提供更有效的风险管理工具。"),
    ],
    "国外科技": [
        ("Nvidia 发布 NemoClaw 安全平台", "Nvidia 在 GTC 2026 上推出 NemoClaw，为 OpenClaw 提供隔离沙盒环境，增强隐私和安全保护。"),
        ("Dell Precision 笔记本回归", "Dell 发布新款 Pro Precision 5 和 7 系列，搭载 Intel Panther Lake 处理器和 Nvidia Blackwell GPU。"),
        ("苹果收购 MotionVFX", "苹果收购 Final Cut Pro 插件开发商 MotionVFX，强化视频编辑 AI 功能，包括 AI 超分辨率和字幕。"),
        ("Nvidia DLSS 5 技术亮相", "Nvidia 发布 DLSS 5 实时神经渲染技术，支持光线追踪和材质增强，今秋登陆多家游戏厂商作品。"),
        ("Phase 鼠标分体变游戏手柄", "Pixelpaw Labs 推出 Phase 鼠标，可对半分体成为双无线游戏控制器，支持 16K DPI 和动作感应。"),
        ("WhatsApp 状态更新将移至聊天标签", "WhatsApp 测试新功能，将状态更新显示在聊天标签顶部，类似 Messenger 和 Instagram 的故事功能。"),
        ("特斯拉 Influencer 离开粉丝圈", "部分特斯拉忠实支持者因马斯克政治立场和 FSD 过度宣传而转向，社群出现分歧。"),
        ("量子计算新突破", "IBM 宣布量子计算机实现 1000 量子位稳定运行，错误率降低 40%，商业化进程加速。"),
        ("SpaceX 星舰第四次试飞成功", "SpaceX 星舰完成首次轨道级回收，为火星任务奠定基础，马斯克称 2030 年前实现载人登陆。"),
        ("TikTok 算法透明度提升", "TikTok 发布算法透明度报告，允许用户查看推荐系统工作原理，回应监管压力。"),
    ],
    "AI 前沿": [
        ("Anthropic 起诉国防部 AI 合同", "Anthropic 对国防部 AI 合同提起诉讼，质疑 AI 在军事决策中的应用边界和伦理问题。"),
        ("AI 模型被用于诈骗", "WIRED 调查发现，数十个 Telegram 频道招募 AI 面容模型，用于创建虚假视频诈骗受害者。"),
        ("Palantir 演示军事 AI 计划生成", "Palantir 展示如何使用 AI 聊天机器人辅助军方生成作战计划，引发自动化战争伦理讨论。"),
        ("Google Gemini 或将引入广告", "Google 高管表示不排除在 Gemini 中加入广告，探索 AI 服务商业化路径。"),
        ("COBOL 语言仍在使用", "尽管已诞生 60 年，COBOL 仍支撑着数万亿美元的金融交易，被称为编程语言的石棉。"),
        ("日本批准首个人体细胞重编程疗法", "日本成为全球首个批准诱导多能干细胞治疗的国家，用于治疗脊髓损伤，开启再生医学新纪元。"),
        ("Grammarly 因 AI 功能遭集体诉讼", "用户起诉 Grammarly 的 AI 专家审查功能误导消费者，称其并非真正的人工专家审查。"),
        ("中国 OpenClaw 热潮兴起", "中国 AI 公司纷纷投入 OpenClaw 生态，形成淘金热式竞争，市场规模预计突破千亿元。"),
        ("AI 音乐视频生成技术突破", "Nvidia GTC 上展示 AI 生成 keynote 总结音乐视频，AI 内容创作能力持续增强。"),
        ("游戏行业 AI 担忧成真", "游戏开发者对 AI 的担忧正在成为现实，包括版权争议和就业岗位流失问题。"),
    ],
    "加密货币": [
        ("比特币价格波动加剧", "比特币在 68,000-72,000 美元区间震荡，市场等待新的催化剂，机构持仓保持稳定。"),
        ("以太坊创始人谈节点简化", "Vitalik Buterin 呼吁简化以太坊节点运行流程，降低用户自主验证门槛。"),
        ("SEC 批准更多 ETF 申请", "SEC 继续批准加密货币 ETF 产品，传统金融机构入场加速，市场流动性提升。"),
        ("稳定币监管框架提案", "美国国会提出稳定币监管框架，要求发行商保持充足储备并定期审计。"),
        ("NFT 市场复苏迹象", "NFT 交易量环比增长 35%，蓝筹项目领头的复苏可能预示市场情绪好转。"),
        ("Layer2 解决方案竞争白热化", "各大 Layer2 项目争夺市场份额，交易费用降低 90%，用户体验显著提升。"),
        ("DeFi 协议安全性提升", "多 DeFi 协议引入形式化验证，智能合约安全漏洞显著减少，保险覆盖增加。"),
        ("央行数字货币进展加速", "全球 130 多国探索 CBDC，中国数字人民币试点扩展，美联储继续研究数字美元。"),
        ("Web3 游戏融资回暖", "Web3 游戏项目获得 5 亿美元融资，游戏公链和基础设施成为投资热点。"),
        ("加密监管清晰度提升", "多国明确加密资产分类和税务处理，行业合规化进程加快。"),
    ]
}

# 添加各栏目内容
for category, news_list in news_data.items():
    # 栏目标题
    heading = doc.add_heading(category, level=1)
    
    # 新闻列表
    for title_text, summary in news_list:
        # 新闻标题
        p = doc.add_paragraph()
        title_run = p.add_run(f"• {title_text}")
        title_run.bold = True
        title_run.font.size = Pt(11)
        
        # 新闻摘要
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.left_indent = Pt(20)
        summary_para.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()  # 栏目间空行

# 添加文档末尾署名
doc.add_paragraph()
doc.add_paragraph()
signature = doc.add_paragraph("🐕 二狗早间报道 · 数据源自公开渠道")
signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
signature.runs[0].italic = True

# 保存文档
output_path = "/mnt/d/个人文件/Desktop/早间报道_20260317.docx"
doc.save(output_path)

print(f"文档已保存至：{output_path}")
