#!/usr/bin/env python3
"""
RSS 新闻聚合日报生成器
使用 RSS 订阅源获取真实新闻数据
"""

import os
import re
import ssl
import json
import gzip
import random
import xml.etree.ElementTree as ET
from datetime import datetime
from urllib.request import Request, urlopen
from urllib.error import URLError
from typing import List, Dict, Tuple
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 禁用 SSL 验证（某些 RSS 源证书问题）
ssl_context = ssl.create_default_context()
ssl_context.check_hostname = False
ssl_context.verify_mode = ssl.CERT_NONE


@dataclass
class NewsItem:
    """新闻条目"""
    title: str
    summary: str
    link: str
    source: str
    pub_date: str = ""


class RSSFetcher:
    """RSS 抓取器"""
    
    # RSS 源配置
    RSS_SOURCES = {
        "国外财经": [
            ("Reuters Business", "https://www.reutersagency.com/feed/?taxonomy=markets&post_type=reuters-best"),
            ("CNBC Finance", "https://www.cnbc.com/id/10000664/device/rss/rss.html"),
            ("MarketWatch", "https://feeds.marketwatch.com/marketwatch/topstories/"),
        ],
        "国外科技": [
            ("TechCrunch", "https://techcrunch.com/feed/"),
            ("The Verge", "https://www.theverge.com/rss/index.xml"),
            ("Ars Technica", "https://feeds.arstechnica.com/arstechnica/index"),
            ("Wired", "https://www.wired.com/feed/rss"),
        ],
        "AI 前沿": [
            ("MIT Tech Review", "https://www.technologyreview.com/feed/"),
            ("VentureBeat AI", "https://venturebeat.com/category/ai/feed/"),
            ("AI Trends", "https://www.aitrends.com/feed/"),
        ],
        "加密货币": [
            ("CoinDesk", "https://www.coindesk.com/arc/outboundfeeds/rss/"),
            ("Cointelegraph", "https://cointelegraph.com/rss"),
            ("CryptoNews", "https://cryptonews.com/news/feed"),
        ],
        "国内新闻": [
            ("新华社", "http://www.xinhuanet.com/rss/xinhua_news.xml"),
            ("人民网", "http://rss.people.com.cn/"),
            ("财新", "https://www.caixin.com/search/rss.xml"),
        ],
        "国外新闻": [
            ("BBC News", "http://feeds.bbci.co.uk/news/rss.xml"),
            ("Reuters World", "https://www.reutersagency.com/feed/?taxonomy=markets&post_type=reuters-best"),
            ("AP News", "https://apnews.com/rss"),
        ],
        "科技前沿": [
            ("Science", "https://www.science.org/rss/news_current.xml"),
            ("Nature News", "https://www.nature.com/news.rss"),
            ("Space.com", "https://www.space.com/feeds/all"),
        ],
        "财经加密": [
            ("Bloomberg", "https://feeds.bloomberg.com/markets/news.rss"),
            ("WSJ Finance", "https://feeds.a.dj.com/rss/RSSMarketsMain.xml"),
            ("CoinDesk Markets", "https://www.coindesk.com/arc/outboundfeeds/rss/?outputType=xml"),
        ]
    }
    
    # User-Agent 列表，轮换使用
    USER_AGENTS = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    ]
    
    def __init__(self):
        self.success_count = 0
        self.fail_count = 0
    
    def fetch_rss(self, url: str, source_name: str) -> List[NewsItem]:
        """抓取单个 RSS 源"""
        try:
            headers = {
                "User-Agent": random.choice(self.USER_AGENTS),
                "Accept": "application/rss+xml, application/xml, text/xml, */*",
                "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8,zh;q=0.7",
                "Accept-Encoding": "identity",  # 禁用 gzip，避免解码问题
                "Connection": "keep-alive",
            }
            
            req = Request(url, headers=headers)
            
            with urlopen(req, context=ssl_context, timeout=15) as response:
                content = response.read()
                
                # 检查是否是 gzip 压缩
                if content[:2] == b'\x1f\x8b':  # gzip magic number
                    try:
                        content = gzip.decompress(content)
                    except:
                        pass
                
                # 尝试解码
                try:
                    xml_content = content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        xml_content = content.decode('gbk')
                    except:
                        xml_content = content.decode('latin-1')
            
            # 解析 XML
            root = ET.fromstring(xml_content)
            
            # 查找 item 元素
            items = []
            for item in root.findall('.//item')[:15]:  # 取前15条
                title = item.find('title')
                desc = item.find('description')
                link = item.find('link')
                pub_date = item.find('pubDate')
                
                title_text = title.text if title is not None else "无标题"
                desc_text = desc.text if desc is not None else ""
                link_text = link.text if link is not None else ""
                date_text = pub_date.text if pub_date is not None else ""
                
                # 清理 HTML 标签
                desc_text = self._clean_html(desc_text)
                
                if title_text and len(title_text) > 5:  # 过滤无效标题
                    items.append(NewsItem(
                        title=title_text.strip(),
                        summary=desc_text[:100] + "..." if len(desc_text) > 100 else desc_text,
                        link=link_text.strip() if link_text else "",
                        source=source_name,
                        pub_date=date_text[:16] if date_text else ""
                    ))
            
            self.success_count += 1
            print(f"  ✅ {source_name}: {len(items)} 条新闻")
            return items
            
        except Exception as e:
            self.fail_count += 1
            print(f"  ❌ {source_name}: {str(e)[:50]}")
            return []
    
    def _clean_html(self, text: str) -> str:
        """清理 HTML 标签和实体"""
        if not text:
            return ""
        # 移除 HTML 标签
        text = re.sub(r'<[^>]+>', '', text)
        # 解码常见 HTML 实体
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        text = text.replace('&#39;', "'")
        # 移除多余空白
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def fetch_category(self, category: str) -> List[NewsItem]:
        """抓取某个栏目的所有 RSS 源"""
        if category not in self.RSS_SOURCES:
            return []
        
        print(f"\n📰 正在抓取栏目: {category}")
        all_news = []
        sources = self.RSS_SOURCES[category]
        
        for source_name, url in sources:
            news_items = self.fetch_rss(url, source_name)
            all_news.extend(news_items)
        
        # 去重（基于标题）
        seen_titles = set()
        unique_news = []
        for item in all_news:
            # 简化标题用于去重比较
            simplified = re.sub(r'[^\w]', '', item.title.lower())
            if simplified and simplified not in seen_titles:
                seen_titles.add(simplified)
                unique_news.append(item)
        
        # 随机打乱，增加多样性
        random.shuffle(unique_news)
        
        return unique_news[:10]  # 每个栏目最多10条


class RSSNewsReportGenerator:
    """RSS 新闻日报生成器"""
    
    MORNING_CATEGORIES = ["国外财经", "国外科技", "AI 前沿", "加密货币"]
    EVENING_CATEGORIES = ["国内新闻", "国外新闻", "科技前沿", "财经加密"]
    
    def __init__(self):
        self.fetcher = RSSFetcher()
        self.today = datetime.now()
        self.date_str = self.today.strftime("%Y年%m月%d日")
        self.date_file = self.today.strftime("%Y%m%d")
    
    def generate_report(self, report_type: str = "morning", 
                       output_dir: str = "/mnt/d/个人文件/Desktop") -> str:
        """生成日报"""
        categories = self.MORNING_CATEGORIES if report_type == "morning" else self.EVENING_CATEGORIES
        title_text = "早间报道" if report_type == "morning" else "晚间新闻报道"
        
        print(f"\n{'='*60}")
        print(f"🚀 开始生成 {title_text} - {self.date_str}")
        print(f"{'='*60}")
        
        # 抓取所有栏目新闻
        all_news = {}
        for category in categories:
            news_items = self.fetcher.fetch_category(category)
            all_news[category] = news_items
        
        # 打印统计
        print(f"\n📊 抓取统计:")
        print(f"  - 成功: {self.fetcher.success_count} 个 RSS 源")
        print(f"  - 失败: {self.fetcher.fail_count} 个 RSS 源")
        total_news = sum(len(items) for items in all_news.values())
        print(f"  - 总计: {total_news} 条新闻")
        
        # 生成文档
        output_path = self._create_word_doc(title_text, all_news, output_dir)
        return output_path
    
    def _create_word_doc(self, title: str, news_data: Dict[str, List[NewsItem]], 
                         output_dir: str) -> str:
        """创建 Word 文档"""
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.name = '微软雅黑'
        title_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        title_para.runs[0].font.size = Pt(24)
        title_para.runs[0].font.bold = True
        
        # 日期
        date_para = doc.add_paragraph(self.date_str)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.name = '微软雅黑'
        date_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        date_para.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 栏目内容
        for category, news_list in news_data.items():
            # 栏目标题
            section_title = doc.add_heading(category, level=1)
            section_title.runs[0].font.name = '微软雅黑'
            section_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            section_title.runs[0].font.size = Pt(16)
            section_title.runs[0].font.bold = True
            
            if not news_list:
                para = doc.add_paragraph("暂无重要新闻（RSS 源暂时无法访问）")
                para.runs[0].font.name = '微软雅黑'
                para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                para.runs[0].italic = True
            else:
                for i, news in enumerate(news_list, 1):
                    # 新闻标题
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(f"{i}. {news.title}")
                    title_run.font.bold = True
                    title_run.font.name = '微软雅黑'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    title_run.font.size = Pt(12)
                    
                    # 新闻摘要和来源
                    summary_para = doc.add_paragraph()
                    summary_text = f"   {news.summary}"
                    if news.source:
                        summary_text += f" [{news.source}]"
                    summary_run = summary_para.add_run(summary_text)
                    summary_run.font.name = '微软雅黑'
                    summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    summary_run.font.size = Pt(10)
            
            doc.add_paragraph()
        
        # 分隔线
        doc.add_paragraph('─' * 50)
        
        # 署名
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = signature.add_run(f"🐕 二狗{title} · RSS 聚合")
        sig_run.font.name = '微软雅黑'
        sig_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        sig_run.font.size = Pt(10)
        
        # 保存
        report_type_suffix = "早间报道" if "早间" in title else "晚间报道"
        output_path = os.path.join(output_dir, f"{report_type_suffix}_{self.date_file}.docx")
        
        os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)
        
        print(f"\n✅ 文档已保存: {output_path}")
        return output_path


def main():
    """主程序"""
    import argparse
    
    parser = argparse.ArgumentParser(description='RSS 新闻日报生成器')
    parser.add_argument('--type', choices=['morning', 'evening'], default='morning',
                       help='报告类型: morning(早间) 或 evening(晚间)')
    parser.add_argument('--output', default='/mnt/d/个人文件/Desktop',
                       help='输出目录')
    
    args = parser.parse_args()
    
    generator = RSSNewsReportGenerator()
    
    try:
        output_path = generator.generate_report(args.type, args.output)
        print(f"\n🎉 日报生成成功！")
        print(f"📄 文件路径: {output_path}")
    except Exception as e:
        print(f"\n❌ 生成失败: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
