"""
实时新闻抓取与日报生成器
使用 Brave Search API 获取真实新闻数据
"""

import os
import re
import json
import time
import logging
from datetime import datetime
from typing import List, Dict, Optional
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Brave API 配置
BRAVE_API_KEY = "BSADt6TZ_kQ3xTcCCDSheOSUnJfQwLm"
BRAVE_API_URL = "https://api.search.brave.com/res/v1/news/search"

# 翻译服务配置
TRANSLATE_API = "https://translate.googleapis.com/translate_a/single"


class Translator:
    """中英文翻译器（使用 Google Translate 免费接口）"""
    
    def __init__(self):
        self.cache = {}  # 翻译缓存，避免重复翻译
    
    def translate_to_chinese(self, text: str) -> str:
        """将英文翻译成中文"""
        if not text or not text.strip():
            return text
        
        # 检查缓存
        if text in self.cache:
            return self.cache[text]
        
        # 如果文本主要是中文，直接返回
        chinese_ratio = sum(1 for c in text if '\u4e00' <= c <= '\u9fff') / len(text)
        if chinese_ratio > 0.3:
            return text
        
        try:
            import subprocess
            import urllib.parse
            
            # 使用 curl 调用 Google Translate 免费接口
            params = {
                "client": "gtx",
                "sl": "en",
                "tl": "zh-CN",
                "dt": "t",
                "q": text
            }
            url = f"{TRANSLATE_API}?{urllib.parse.urlencode(params)}"
            
            cmd = ["curl", "-s", "-L", "--connect-timeout", "10", "--max-time", "30", url]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=35)
            
            if result.returncode == 0 and result.stdout:
                # 解析返回结果 [[["translated_text","original_text",...],...],...]
                import json
                data = json.loads(result.stdout)
                if data and data[0]:
                    translated = "".join([part[0] for part in data[0] if part[0]])
                    self.cache[text] = translated
                    return translated
        except Exception as e:
            logger.warning(f"翻译失败: {e}, 保留原文")
        
        return text
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """批量翻译"""
        return [self.translate_to_chinese(t) for t in texts]

@dataclass
class NewsItem:
    """新闻条目数据结构"""
    title: str
    summary: str
    url: str
    source: str
    published_date: str

class BraveNewsFetcher:
    """Brave Search API 新闻抓取器"""
    
    def __init__(self, api_key: str, translator: Translator = None):
        self.api_key = api_key
        self.headers = {
            "Accept": "application/json",
            "X-Subscription-Token": api_key
        }
        self.translator = translator
    
    def fetch_news(self, query: str, count: int = 10, freshness: str = "day") -> List[NewsItem]:
        """抓取新闻 - 使用 curl 命令避免 Python SSL 问题"""
        try:
            import subprocess
            import urllib.parse
            import json as json_lib
            
            params = {
                "q": query,
                "count": min(count, 50),
                "freshness": freshness,
                "safesearch": "off"
            }
            
            logger.info(f"Fetching news for query: {query}")
            
            # 构建 URL
            query_string = urllib.parse.urlencode(params)
            url = f"{BRAVE_API_URL}?{query_string}"
            
            # 使用 curl 命令
            cmd = [
                "curl", "-s", "-L",
                "-H", f"Accept: {self.headers['Accept']}",
                "-H", f"X-Subscription-Token: {self.headers['X-Subscription-Token']}",
                "--connect-timeout", "30",
                "--max-time", "60",
                url
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=65)
            
            if result.returncode != 0:
                logger.error(f"curl error: {result.stderr}")
                return []
            
            data = json_lib.loads(result.stdout)
            news_items = []
            
            for item in data.get("results", []):
                title = item.get("title", "无标题")
                summary = self._clean_summary(item.get("description", ""))
                
                # 翻译标题和摘要
                if self.translator:
                    title = self.translator.translate_to_chinese(title)
                    summary = self.translator.translate_to_chinese(summary)
                    time.sleep(0.1)  # 避免翻译API限流
                
                news_item = NewsItem(
                    title=title,
                    summary=summary,
                    url=item.get("url", ""),
                    source=self._extract_source(item.get("url", "")),
                    published_date=item.get("pub_date", "")[:10] if item.get("pub_date") else ""
                )
                news_items.append(news_item)
            
            logger.info(f"Fetched {len(news_items)} news items for {query}")
            return news_items
            
        except Exception as e:
            logger.error(f"Fetch error: {e}")
            return []
    
    def _clean_summary(self, text: str, max_length: int = 80) -> str:
        """清理摘要文本"""
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) > max_length:
            text = text[:max_length-3] + "..."
        return text if text else "暂无摘要"
    
    def _extract_source(self, url: str) -> str:
        """从 URL 提取来源"""
        try:
            from urllib.parse import urlparse
            domain = urlparse(url).netloc
            return domain.replace("www.", "")
        except:
            return "未知来源"


class DailyNewsReportGenerator:
    """日报生成器"""
    
    # 早间报道配置
    MORNING_CATEGORIES = {
        "国外财经": [
            "global stock market today S&P 500 Nasdaq",
            "Federal Reserve interest rate decision today",
            "oil price Brent crude today",
            "gold price today record high"
        ],
        "国外科技": [
            "tech news today The Verge",
            "Apple iPhone news today",
            "Nvidia AI chip news",
            "Samsung Galaxy news"
        ],
        "AI 前沿": [
            "OpenAI GPT news today",
            "Google Gemini AI news",
            "Anthropic Claude new features",
            "AI breakthrough 2026"
        ],
        "加密货币": [
            "Bitcoin price today",
            "Ethereum news today",
            "cryptocurrency market news",
            "CoinDesk latest crypto"
        ]
    }
    
    # 晚间报道配置
    EVENING_CATEGORIES = {
        "国内新闻": [
            "China news today Xinhua",
            "China economic policy news",
            "China tech news today"
        ],
        "国外新闻": [
            "world news today BBC",
            "Reuters world news today",
            "CNN news today",
            "international news today"
        ],
        "科技前沿": [
            "technology breakthrough news today",
            "space exploration news today",
            "quantum computing news",
            "biotech news today"
        ],
        "财经加密": [
            "stock market closing today",
            "forex news today",
            "crypto market today",
            "DeFi news today"
        ]
    }
    
    def __init__(self, api_key: str):
        self.translator = Translator()
        self.fetcher = BraveNewsFetcher(api_key, self.translator)
    
    def generate_report(self, report_type: str = "morning", output_dir: str = "/mnt/d/个人文件/Desktop") -> str:
        """
        生成日报
        
        Args:
            report_type: "morning" 或 "evening"
            output_dir: 输出目录
        
        Returns:
            生成的文件路径
        """
        today = datetime.now()
        date_str = today.strftime("%Y年%m月%d日")
        date_file = today.strftime("%Y%m%d")
        
        # 选择分类
        categories = self.MORNING_CATEGORIES if report_type == "morning" else self.EVENING_CATEGORIES
        title_text = "早间报道" if report_type == "morning" else "晚间新闻报道"
        
        logger.info(f"开始生成 {title_text}...")
        
        # 抓取新闻
        all_news = {}
        for category, queries in categories.items():
            logger.info(f"抓取栏目: {category}")
            category_news = []
            
            for query in queries:
                news_items = self.fetcher.fetch_news(query, count=5, freshness="day")
                category_news.extend(news_items)
                time.sleep(0.5)  # 避免请求过快
            
            # 去重并限制数量
            seen_titles = set()
            unique_news = []
            for item in category_news:
                if item.title not in seen_titles and len(unique_news) < 10:
                    seen_titles.add(item.title)
                    unique_news.append(item)
            
            all_news[category] = unique_news
        
        # 生成 Word 文档
        return self._create_word_doc(title_text, date_str, date_file, all_news, output_dir)
    
    def _create_word_doc(self, title: str, date_str: str, date_file: str, 
                         news_data: Dict, output_dir: str) -> str:
        """创建 Word 文档"""
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.name = '微软雅黑'
        title_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        title_para.runs[0].font.size = Pt(24)
        title_para.runs[0].font.bold = True
        
        # 日期
        date_para = doc.add_paragraph(date_str)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.name = '微软雅黑'
        date_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        date_para.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 栏目内容
        for category, news_list in news_data.items():
            # 栏目标题
            section_title = doc.add_heading(category, level=1)
            section_title.runs[0].font.name = '微软雅黑'
            section_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            section_title.runs[0].font.size = Pt(16)
            section_title.runs[0].font.bold = True
            
            if not news_list:
                para = doc.add_paragraph("暂无重要新闻")
                para.runs[0].font.name = '微软雅黑'
                para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                for i, news in enumerate(news_list, 1):
                    # 新闻标题
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(f"{i}. {news.title}")
                    title_run.font.bold = True
                    title_run.font.name = '微软雅黑'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    title_run.font.size = Pt(12)
                    
                    # 新闻摘要和来源
                    summary_para = doc.add_paragraph()
                    summary_text = f"   {news.summary}"
                    if news.source:
                        summary_text += f" [{news.source}]"
                    summary_run = summary_para.add_run(summary_text)
                    summary_run.font.name = '微软雅黑'
                    summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    summary_run.font.size = Pt(10)
            
            doc.add_paragraph()
        
        # 分隔线
        doc.add_paragraph('─' * 50)
        
        # 署名
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = signature.add_run(f"🐕 二狗{title} · 数据源自 Brave Search")
        sig_run.font.name = '微软雅黑'
        sig_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        sig_run.font.size = Pt(10)
        
        # 保存
        report_type_suffix = "早间报道" if "早间" in title else "晚间报道"
        output_path = os.path.join(output_dir, f"{report_type_suffix}_{date_file}.docx")
        
        os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"文档已保存: {output_path}")
        return output_path


def main():
    """主程序"""
    import argparse
    
    parser = argparse.ArgumentParser(description='实时新闻日报生成器')
    parser.add_argument('--type', choices=['morning', 'evening'], default='morning',
                       help='报告类型: morning(早间) 或 evening(晚间)')
    parser.add_argument('--output', default='/mnt/d/个人文件/Desktop',
                       help='输出目录')
    
    args = parser.parse_args()
    
    # 检查依赖
    try:
        import requests
    except ImportError:
        print("❌ 请先安装依赖: pip install requests python-docx")
        return
    
    # 生成报告
    generator = DailyNewsReportGenerator(BRAVE_API_KEY)
    
    try:
        output_path = generator.generate_report(args.type, args.output)
        print(f"\n✅ 日报生成成功！")
        print(f"📄 文件路径: {output_path}")
    except Exception as e:
        logger.error(f"生成失败: {e}")
        print(f"\n❌ 生成失败: {e}")


if __name__ == "__main__":
    main()
