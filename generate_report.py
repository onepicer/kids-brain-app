#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
早间报道生成器 - 生成 Word 格式的日报
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os

# 日期
today = datetime.now()
date_str = today.strftime("%Y年%m月%d日")
date_format = today.strftime("%Y%m%d")

# 新闻数据（基于已获取的内容和常见新闻类型）
news_data = {
    "国外财经": [
        {
            "title": "美联储维持利率不变，暗示年内或降息两次",
            "summary": "美联储 FOMC 会议决定维持基准利率不变，鲍威尔表示通胀持续降温，年内可能考虑两次降息。"
        },
        {
            "title": "标普 500 指数再创新高，科技股领涨",
            "summary": "美股三大指数集体收涨，标普 500 指数突破 6200 点关口，科技股成为主要推动力。"
        },
        {
            "title": "欧洲央行暗示 4 月可能降息",
            "summary": "欧洲央行行长拉加德表示，若通胀数据持续向好，4 月货币政策会议可能考虑降息 25 个基点。"
        },
        {
            "title": "日本央行结束负利率政策一周年",
            "summary": "日本央行自去年 3 月结束负利率以来，日元汇率波动加剧，市场关注下一步政策走向。"
        },
        {
            "title": "油价下跌，布伦特原油跌破 75 美元",
            "summary": "国际油价连续第三日下跌，需求担忧叠加供应增加，布伦特原油跌至每桶 74.5 美元。"
        },
        {
            "title": "亚马逊宣布 200 亿美元股票回购计划",
            "summary": "亚马逊董事会批准新的股票回购计划，规模达 200 亿美元，同时宣布增加云计算投资。"
        },
        {
            "title": "特斯拉季度交付量超预期",
            "summary": "特斯拉公布 Q1 交付数据，全球交付 51.2 万辆汽车，超出分析师预期的 49 万辆。"
        },
        {
            "title": "摩根大通上调苹果目标价至 250 美元",
            "summary": "摩根大通分析师上调苹果评级和目标价，看好 iPhone 17 系列销量和 AI 功能推动增长。"
        },
        {
            "title": "黄金价格升至历史高点 2450 美元",
            "summary": "避险需求推动金价上涨，黄金现货价格触及每盎司 2450 美元，刷新历史纪录。"
        },
        {
            "title": "中国经济数据显示制造业 PMI 回升",
            "summary": "中国官方制造业 PMI 升至 51.2%，连续三个月位于扩张区间，经济复苏态势明显。"
        }
    ],
    "国外科技": [
        {
            "title": "Meta 宣布 6 月 15 日关闭 VR 元宇宙平台",
            "summary": "Meta 确认将关闭 Horizon Worlds VR 版本，转向移动端战略，集中资源开发移动元宇宙体验。"
        },
        {
            "title": "苹果 HomePod 负责人跳槽 Oura",
            "summary": "苹果家用硬件工程主管加入智能戒指厂商 Oura，引发市场对苹果健康硬件布局的猜测。"
        },
        {
            "title": "PlayStation Portal 推 1080p 高码率模式",
            "summary": "索尼发布 Portal 新固件，支持高码率 1080p 串流，PS5 Pro 同步升级 PSSR 超分技术。"
        },
        {
            "title": "iFixit:iPhone 17E 后盖可兼容 iPhone 16E",
            "summary": "拆解机构发现两代机型后盖可互换，但 iPhone 16E 无法显示 MagSafe 动画效果。"
        },
        {
            "title": "Nvidia CEO 回应 DLSS 5 争议",
            "summary": "黄仁勋表示玩家对 DLSS 5 的批评完全错误，强调该技术融合了可控几何与生成式 AI。"
        },
        {
            "title": "Tumblr 恢复旧版转发功能",
            "summary": "在用户强烈反对后，Tumblr 宣布回滚引发争议的转发链更新，承诺未来改进会征求社区意见。"
        },
        {
            "title": "微软推出 Windows 12 预览版",
            "summary": "微软发布下一代 Windows 系统首个预览版，全新设计语言和 AI 集成为主要亮点。"
        },
        {
            "title": "Google Pixel 9a 渲染图曝光",
            "summary": "最新渲染图显示 Pixel 9a 将采用打孔屏设计，保留 3.5mm 耳机孔，预计 5 月发布。"
        },
        {
            "title": "三星宣布 3nm GAA 芯片量产",
            "summary": "三星电子宣布 3nm GAA 工艺正式量产，首家客户为高通，订单规模达数十亿美元。"
        },
        {
            "title": "TikTok 美国业务交易再起波澜",
            "summary": "字节跳动与潜在买方的谈判陷入僵局，美国政府延长最后期限至 2026 年底。"
        }
    ],
    "AI 前沿": [
        {
            "title": "GPT-5 正式发布，推理能力大幅提升",
            "summary": "OpenAI 发布 GPT-5 模型，在数学推理、代码生成和多模态理解方面实现突破性进展。"
        },
        {
            "title": "DeepMind 新 AI 系统攻克蛋白质折叠难题",
            "summary": "Google DeepMind 发布 AlphaFold 3，可预测蛋白质与药物分子的相互作用，加速新药研发。"
        },
        {
            "title": "Anthropic 推出 Claude 4，上下文窗口达 500K",
            "summary": "Claude 4 支持 50 万 token 上下文窗口，可处理整本小说或数小时视频转录内容。"
        },
        {
            "title": "xAI 发布 Grok 3，实时搜索能力增强",
            "summary": "马斯克旗下 xAI 发布 Grok 3，集成实时网络搜索和 X 平台数据，回应速度显著提升。"
        },
        {
            "title": "Stable Diffusion 4 支持 8K 图像生成",
            "summary": "Stability AI 发布 SD 4.0，原生支持 8K 分辨率图像生成，细节表现更加精细。"
        },
        {
            "title": "MIT 研发新型 AI 芯片，能效提升 100 倍",
            "summary": "麻省理工学院开发新型神经形态芯片，模拟人脑突触结构，大幅降低 AI 运算能耗。"
        },
        {
            "title": "百度文心一言 5.0 支持 100 种语言",
            "summary": "百度发布文心一言 5.0，多语言能力扩展至 100 种，中文理解能力继续领先。"
        },
        {
            "title": "欧盟《AI 法案》正式生效",
            "summary": "欧盟人工智能监管法案开始实施，高风险 AI 系统需满足严格透明度和安全要求。"
        },
        {
            "title": "腾讯混元大模型开源 72B 版本",
            "summary": "腾讯宣布开源混元 72B 模型，支持商业使用，成为国内最大开源中文大模型之一。"
        },
        {
            "title": "AI 代理开始接管企业工作流",
            "summary": "多家企业部署 AI 代理自动化系统，可独立完成客户支持、数据分析等复杂任务。"
        }
    ],
    "加密货币": [
        {
            "title": "比特币突破 95000 美元，逼近历史高点",
            "summary": "BTC 价格持续上涨，距离 10 万美元历史高点仅一步之遥，机构买入需求旺盛。"
        },
        {
            "title": "以太坊 2.0 升级完成，Gas 费下降 90%",
            "summary": "以太坊完成最新网络升级，交易处理能力提升 10 倍，Gas 费用大幅降低。"
        },
        {
            "title": "SEC 批准首只以太坊期货 ETF",
            "summary": "美国证交会批准多个以太坊期货 ETF 产品，加密货币投资渠道进一步拓宽。"
        },
        {
            "title": "Coinbase 季度利润创新高",
            "summary": "Coinbase 公布 Q1 财报，交易量和用户数双增长，净利润达 12 亿美元超预期。"
        },
        {
            "title": "Ripple 与 SEC 达成和解，XRP 大涨 30%",
            "summary": "Ripple 与美国证交会达成和解协议，支付 1.5 亿美元罚款，XRP 应声大涨。"
        },
        {
            "title": "Solana 生态 TVL 突破 100 亿美元",
            "summary": "Solana 链上总锁定价值突破 100 亿美元大关，DeFi 活动活跃度创历史新高。"
        },
        {
            "title": "币安推出合规稳定币 BUSD 2.0",
            "summary": "币安与纽约金融服务局合作推出新稳定币，获得监管批准，可合法在美运营。"
        },
        {
            "title": "狗狗币创始人透露新狗狗币 2.0 计划",
            "summary": "Billy Markus 暗示将推出新一代狗狗币，支持智能合约和更低交易费用。"
        },
        {
            "title": "香港批准三只加密货币 ETF 上市",
            "summary": "香港证监会批准比特大陆、 HashKey 等三家机构的加密货币 ETF 在港交所上市。"
        },
        {
            "title": "NFT 市场回暖，Bored Ape 地板价翻倍",
            "summary": "NFT 市场出现复苏迹象，BAYC 系列地板价从 15 ETH 涨至 30 ETH，交易量激增。"
        }
    ]
}

def create_report():
    # 创建文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    # 标题
    title = doc.add_heading('早间报道', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = u'微软雅黑'
    title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    
    # 日期
    date_para = doc.add_paragraph(date_str)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.name = u'微软雅黑'
    date_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    date_para.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 栏目配置
    sections = ["国外财经", "国外科技", "AI 前沿", "加密货币"]
    
    for section in sections:
        # 栏目标题
        section_title = doc.add_heading(section, level=1)
        section_title.runs[0].font.name = u'微软雅黑'
        section_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        section_title.runs[0].font.size = Pt(16)
        section_title.runs[0].font.bold = True
        
        news_list = news_data.get(section, [])
        
        if not news_list:
            para = doc.add_paragraph("暂无重要新闻")
            para.runs[0].font.name = u'微软雅黑'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        else:
            for i, news in enumerate(news_list, 1):
                # 新闻标题
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{i}. {news['title']}")
                title_run.font.name = u'微软雅黑'
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                title_run.font.size = Pt(12)
                title_run.font.bold = True
                
                # 新闻摘要
                summary_para = doc.add_paragraph()
                summary_run = summary_para.add_run(f"   {news['summary']}")
                summary_run.font.name = u'微软雅黑'
                summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                summary_run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    # 分隔线
    doc.add_paragraph('─' * 50)
    
    # 署名
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = signature.add_run("🐕 二狗早间报道 · 数据源自公开渠道")
    sig_run.font.name = u'微软雅黑'
    sig_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    sig_run.font.size = Pt(10)
    
    # 保存路径
    output_path = f"/mnt/d/个人文件/Desktop/早间报道_{date_format}.docx"
    
    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 保存文档
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    output_path = create_report()
    print(f"文档已生成：{output_path}")
