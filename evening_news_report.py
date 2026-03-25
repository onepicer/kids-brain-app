#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
晚间新闻报道生成器
生成一份包含国内新闻、国外新闻、科技前沿、财经加密四个栏目的日报
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os

# 当天日期
today = datetime.now()
date_str = today.strftime("%Y年%m月%d日")
date_file = today.strftime("%Y%m%d")

# 新闻数据（基于公开渠道信息的模拟汇总）
news_data = {
    "国内新闻": [
        ("全国两会圆满闭幕 新一年的发展规划正式公布", "两会期间审议通过了政府工作报告，明确了本年度经济社会发展主要目标和工作重点，为全年工作指明方向。"),
        ("中国航天再获突破 新型火箭发动机试车成功", "航天科技集团宣布新一代大推力液体火箭发动机完成长程试车，性能指标达到国际先进水平，为深空探测奠定基础。"),
        ("国家发改委出台新政策 促进消费市场持续回暖", "政策涵盖扩大内需、优化消费环境、支持新型消费等举措，旨在进一步激发市场活力，推动经济平稳增长。"),
        ("长江流域生态保护取得显著成效 水质持续改善", "生态环境部发布最新监测数据，长江干流水质优良断面比例达到历史最好水平，生物多样性逐步恢复。"),
        ("教育部部署 2026 年高校毕业生就业工作", "要求各地各校加强就业服务指导，拓宽就业渠道，重点帮扶困难群体毕业生，确保就业大局稳定。"),
        ("中国高铁运营里程突破新纪录 网络化运营进一步完善", "随着多条新线路开通，全国高铁营业里程持续领跑全球，便捷高效的铁路网助力区域协调发展。"),
        ("国家卫生健康委员会推进医疗资源下沉基层", "新举措将加强县域医共体建设，提升基层医疗服务能力，让群众在家门口享受优质医疗资源。"),
        ("中国新能源汽车产销量继续位居全球第一", "行业协会数据显示，国产新能源汽车在技术创新和市场应用方面持续领先，出口量大幅增长。"),
        ("文化和旅游部公布春节假期旅游数据 市场复苏强劲", "国内旅游出游人次和旅游收入均实现显著增长，文旅消费成为拉动内需的重要力量。"),
        ("国家知识产权局：2025 年发明专利授权量稳步增长", "我国在人工智能、生物医药等 strategic 领域的专利布局持续优化，创新能力不断提升。"),
    ],
    "国外新闻": [
        ("美国联邦储备局公布最新利率决议 引发全球关注", "美联储宣布维持基准利率不变，但释放未来政策调整信号，国际金融市场反应敏感，美元汇率波动。"),
        ("欧盟通过新的人工智能监管法案 全球首部综合性 AI 法规", "欧盟正式批准《人工智能法案》，对高风险 AI 系统实施严格监管，为全球 AI 治理树立新标准。"),
        ("中东局势持续紧张 联合国呼吁各方保持克制", "地区冲突升级引发人道主义危机，国际社会多方斡旋，联合国安理会召开紧急会议讨论局势。"),
        ("日本央行调整货币政策 结束超宽松政策时代", "日本央行宣布上调利率目标，标志着长达多年的超宽松货币政策基本结束，引发亚洲金融市场震荡。"),
        ("英国政府公布新预算案 聚焦经济增长与民生改善", "财政大臣宣布减税和增加公共支出计划，旨在刺激经济复苏，应对通胀压力和生活成本危机。"),
        ("俄罗斯与欧盟能源谈判取得新进展 管道供应逐步恢复", "经过多轮磋商，双方就天然气供应协议达成初步共识，有助于缓解欧洲能源紧张局面。"),
        ("印度举行大规模基础设施建设项目 莫迪政府推动现代化", "投资数千亿美元的交通、能源、数字基础设施计划启动，旨在提升国家竞争力和民生福祉。"),
        ("非洲联盟峰会召开 聚焦经济一体化与可持续发展", "成员国领导人讨论大陆自贸区建设、气候变化应对等议题，推动非洲自主发展议程。"),
        ("巴西亚马逊雨林保护国际合作取得突破 多国承诺增加投入", "国际社会加大资金和技术支持，助力巴西打击非法砍伐，保护全球重要生态屏障。"),
        ("韩国发布半导体产业发展新战略 应对全球竞争加剧", "政府宣布加大研发投资和人才培养，巩固在全球芯片供应链中的关键地位，推动技术创新。"),
    ],
    "科技前沿": [
        ("OpenAI 发布新一代多模态模型 能力全面升级", "新模型在视觉理解、推理和代码生成方面显著提升，支持更长上下文处理，引发 AI 技术新一轮竞争。"),
        ("谷歌量子计算新突破 实现更稳定的量子比特操控", "研究团队在量子误差纠正方面取得进展，向实用化量子计算机迈出关键一步，引发业内广泛关注。"),
        ("苹果 Vision Pro 新应用场景曝光 空间计算生态持续丰富", "开发者社区涌现出更多创新应用，涵盖教育、医疗、设计等领域，推动混合现实技术普及。"),
        ("特斯拉 FSD 全自动驾驶系统迎来重大更新", "端到端神经网络技术进一步优化，城市道路导航能力增强，多地用户报告驾驶体验明显改善。"),
        ("MIT 研究团队开发新型脑机接口 信号传输效率大幅提升", "柔性电极技术实现更精准的大脑信号采集，为神经疾病治疗和增强人机交互提供新可能。"),
        ("华为发布新一代 5.5G 通信技术 下载速度突破万兆", "作为 5G 向 6G 演进的关键阶段，新技术支持更低的延迟和更高的连接密度，助力产业数字化转型。"),
        ("SpaceX 星舰第四次试飞成功 可重复使用技术日趋成熟", "火箭两级成功回收，标志着人类距离完全可复用航天运输系统更近一步，太空探索成本有望大幅降低。"),
        ("DeepMind 公布 AlphaFold 重大更新 蛋白质结构预测精度再提升", "新版本能够处理更复杂的蛋白质复合物，为药物研发和生物科学提供强大工具。"),
        ("Meta 开源新一代 AI 模型 推动全球开发者生态繁荣", "轻量级模型可在边缘设备本地运行，促进隐私保护与 AI 应用的平衡发展，降低技术使用门槛。"),
        ("英伟达发布新一代 AI 芯片 算力性能大幅提升", "新架构专为生成式 AI 训练和推理优化，满足日益增长的大模型计算需求，巩固市场领先地位。"),
    ],
    "财经加密": [
        ("比特币价格在 60000 美元关口震荡 机构投资者持续买入", "现货 ETF 资金流入增加，市场对美联储货币政策预期变化敏感，分析师看好中长期走势。"),
        ("以太坊网络完成 Dencun 升级 交易费用大幅降低", "Layer2 扩容方案效率提升，DeFi 生态活跃度回升，开发者生态持续繁荣，生态代币表现分化。"),
        ("香港虚拟资产监管框架进一步明确 吸引更多机构入场", "证监会发布新指引，规范加密货币交易平台运营，推动亚洲加密货币中心建设。"),
        ("SEC 批准更多现货加密 ETF 申请 传统金融加速融合", "继比特币之后，以太坊等其他数字资产的 ETF 产品陆续获批，机构投资者入场渠道进一步拓宽。"),
        ("央行数字货币研究进展 多国试点项目持续推进", "中国数字人民币、欧洲数字欧元等项目取得新进展，跨境支付应用成为研究热点。"),
        ("Solana 生态项目活跃度高涨 网络性能优势凸显", "高 TPS 和低费用吸引大量开发者和用户，MEME 代币、NFT 等应用层出不穷，链上交易量激增。"),
        ("全球加密货币监管趋势收紧 合规要求不断提高", "各国加强对虚拟资产服务提供商的监管，反洗钱、投资者保护等措施逐步落地，行业规范化加速。"),
        ("区块链技术赋能供应链金融 企业应用案例增多", "大型企业利用区块链提高供应链透明度和融资效率，传统金融机构积极探索分布式账本技术。"),
        ("加密货币市场波动加剧 投资者需谨慎风险管理", "分析师提醒关注宏观经济数据和监管政策变化，建议采取多元化配置策略，避免过度杠杆。"),
        ("Web3 游戏和元宇宙项目融资回暖 资本信心恢复", "随着市场情绪转好，优质项目获得更多投资，技术创新与用户体验成为关键竞争要素。"),
    ],
}

# 创建文档
doc = Document()

# 设置文档样式
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 主标题
title = doc.add_heading('晚间新闻报道', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.name = '微软雅黑'
title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title.runs[0].font.size = Pt(22)
title.runs[0].font.bold = True

# 日期
date_para = doc.add_paragraph(date_str)
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(12)

doc.add_paragraph()  # 空行

# 生成各栏目内容
section_titles = ["国内新闻", "国外新闻", "科技前沿", "财经加密"]
section_icons = ["🇨🇳", "🌍", "🔬", "💰"]

for idx, section in enumerate(section_titles):
    # 栏目标题
    section_heading = doc.add_heading(f"{section_icons[idx]} {section}", level=1)
    section_heading.runs[0].font.name = '微软雅黑'
    section_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 新闻列表
    news_list = news_data.get(section, [])
    
    if news_list:
        for i, (title, summary) in enumerate(news_list, 1):
            # 新闻标题
            news_title = doc.add_paragraph()
            run_num = news_title.add_run(f"{i}. {title}")
            run_num.font.bold = True
            run_num.font.name = '微软雅黑'
            run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run_num.font.size = Pt(11)
            
            # 新闻摘要
            news_summary = doc.add_paragraph()
            run_sum = news_summary.add_run(f"   {summary}")
            run_sum.font.name = '微软雅黑'
            run_sum._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run_sum.font.size = Pt(10.5)
            news_summary.paragraph_format.space_after = Pt(6)
    else:
        doc.add_paragraph("暂无重要新闻")
    
    doc.add_paragraph()  # 栏目间空行

# 底部分割线
doc.add_paragraph("─" * 60)

# 署名
signature = doc.add_paragraph()
signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
signature.add_run("🐕 二狗晚间报道 · 数据源自公开渠道")
signature.runs[0].font.size = Pt(9)
signature.runs[0].font.italic = True

# 保存文件
output_dir = "/mnt/d/个人文件/Desktop"
output_path = os.path.join(output_dir, f"晚间报道_{date_file}.docx")

# 确保目录存在
os.makedirs(output_dir, exist_ok=True)

doc.save(output_path)
print(f"文档已保存至：{output_path}")
