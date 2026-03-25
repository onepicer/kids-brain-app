#!/usr/bin/env python3
"""
RSS 新闻聚合日报生成器（中文翻译版）
"""

import os
import re
import ssl
import gzip
import random
import xml.etree.ElementTree as ET
from datetime import datetime
from urllib.request import Request, urlopen
from typing import List, Tuple
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ssl_context = ssl.create_default_context()
ssl_context.check_hostname = False
ssl_context.verify_mode = ssl.CERT_NONE


@dataclass
class NewsItem:
    title: str
    title_cn: str
    summary: str
    source: str


# RSS 源（已测试可用）
RSS_SOURCES = {
    "国外新闻": [
        ("NPR News", "https://feeds.npr.org/1004/rss.xml"),
    ],
    "科技新闻": [
        ("TechCrunch", "https://techcrunch.com/feed/"),
        ("Ars Technica", "https://feeds.arstechnica.com/arstechnica/index"),
    ],
    "财经新闻": [
        ("CNBC", "https://www.cnbc.com/id/10000664/device/rss/rss.html"),
    ],
    "加密货币": [
        ("Cointelegraph", "https://cointelegraph.com/rss"),
    ],
}

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
]


def translate_text(text: str) -> str:
    """翻译文本为中文"""
    if not text or len(text.strip()) < 2:
        return text
    
    # 检测是否已经是中文
    if re.match(r'^[\u4e00-\u9fff\s\W\d]+$', text):
        return text
    
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='auto', target='zh-CN')
        result = translator.translate(text[:4500])  # 限制长度
        return result if result else text
    except:
        return text


def fetch_rss(url: str, source_name: str) -> List[Tuple[str, str, str]]:
    """抓取单个 RSS 源，返回 (标题, 摘要, 来源)"""
    items = []
    try:
        headers = {
            "User-Agent": random.choice(USER_AGENTS),
            "Accept": "application/rss+xml, application/xml, text/xml, */*",
            "Accept-Encoding": "identity",
        }
        
        req = Request(url, headers=headers)
        
        with urlopen(req, context=ssl_context, timeout=15) as response:
            content = response.read()
            
            if content[:2] == b'\x1f\x8b':
                content = gzip.decompress(content)
            
            try:
                xml_content = content.decode('utf-8')
            except:
                xml_content = content.decode('latin-1')
        
        root = ET.fromstring(xml_content)
        
        for item in root.findall('.//item')[:8]:  # 每个源最多8条
            title = item.find('title')
            desc = item.find('description')
            
            title_text = title.text if title is not None else ""
            desc_text = desc.text if desc is not None else ""
            
            # 清理 HTML
            desc_text = re.sub(r'<[^>]+>', '', desc_text)
            desc_text = desc_text.replace('&nbsp;', ' ').strip()
            
            if title_text and len(title_text) > 5:
                items.append((
                    title_text.strip(),
                    desc_text[:100] + "..." if len(desc_text) > 100 else desc_text,
                    source_name
                ))
        
        print(f"  ✅ {source_name}: {len(items)} 条")
        
    except Exception as e:
        print(f"  ❌ {source_name}: {str(e)[:40]}")
    
    return items


def generate_report(report_type: str = "morning", output_dir: str = "/mnt/d/个人文件/Desktop") -> str:
    """生成日报"""
    today = datetime.now()
    date_str = today.strftime("%Y年%m月%d日")
    date_file = today.strftime("%Y%m%d")
    
    title_text = "早间报道" if report_type == "morning" else "晚间新闻报道"
    
    print(f"\n{'='*50}")
    print(f"🚀 {title_text} - {date_str}")
    print(f"{'='*50}")
    
    # 抓取新闻
    all_items = []
    
    for category, sources in RSS_SOURCES.items():
        print(f"\n📰 {category}")
        
        for source_name, url in sources:
            items = fetch_rss(url, source_name)
            for title, summary, source in items:
                all_items.append((category, title, summary, source))
    
    # 去重
    seen = set()
    unique_items = []
    for item in all_items:
        key = re.sub(r'[^\w]', '', item[1].lower())
        if key and key not in seen:
            seen.add(key)
            unique_items.append(item)
    
    print(f"\n📊 总计: {len(unique_items)} 条新闻")
    
    # 翻译标题（并行）
    print("\n🔄 正在翻译...")
    titles = [item[1] for item in unique_items]
    
    translated_titles = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(translate_text, t): i for i, t in enumerate(titles)}
        for future in as_completed(futures):
            i = futures[future]
            try:
                translated_titles.append((i, future.result()))
                if len(translated_titles) % 10 == 0:
                    print(f"   📝 已翻译 {len(translated_titles)}/{len(titles)}")
            except:
                translated_titles.append((i, titles[i]))
    
    # 排序
    translated_titles.sort(key=lambda x: x[0])
    cn_titles = [t[1] for t in translated_titles]
    
    print(f"   ✅ 翻译完成")
    
    # 按类别分组
    news_by_category = {}
    for i, (category, title, summary, source) in enumerate(unique_items):
        if category not in news_by_category:
            news_by_category[category] = []
        news_by_category[category].append(NewsItem(
            title=title,
            title_cn=cn_titles[i] if i < len(cn_titles) else title,
            summary=summary,
            source=source
        ))
    
    # 生成 Word
    doc = Document()
    
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title_para = doc.add_heading(title_text, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.size = Pt(24)
    
    # 日期
    date_para = doc.add_paragraph(date_str)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 内容
    for category in RSS_SOURCES.keys():
        news_list = news_by_category.get(category, [])[:6]  # 每类最多6条
        doc.add_heading(category, level=1)
        
        if not news_list:
            doc.add_paragraph("暂无重要新闻")
        else:
            for i, news in enumerate(news_list, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {news.title_cn}").bold = True
                
                # 原文标题（小字）
                if news.title != news.title_cn:
                    op = doc.add_paragraph()
                    op.add_run(f"   [{news.source}] {news.title}").italic = True
        
        doc.add_paragraph()
    
    # 署名
    doc.add_paragraph('─' * 50)
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(f"🐕 二狗{title_text} · RSS 聚合翻译")
    
    # 保存
    filename = "早间报道" if report_type == "morning" else "晚间报道"
    output_path = os.path.join(output_dir, f"{filename}_{date_file}.docx")
    
    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    
    print(f"\n✅ 已保存: {output_path}")
    return output_path


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser()
    parser.add_argument('--type', choices=['morning', 'evening'], default='morning')
    parser.add_argument('--output', default='/mnt/d/个人文件/Desktop')
    
    args = parser.parse_args()
    
    try:
        generate_report(args.type, args.output)
    except Exception as e:
        print(f"❌ 失败: {e}")
        import traceback
        traceback.print_exc()
